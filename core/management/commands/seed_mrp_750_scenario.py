"""Seed (and optionally run + document) the VGS MRP 750 material-planning scenario.

    python manage.py seed_mrp_750_scenario --tenant "VGS and Technologies Pvt Ltd"
    python manage.py seed_mrp_750_scenario --tenant "VGS and Technologies Pvt Ltd" --run
    python manage.py seed_mrp_750_scenario --tenant "VGS and Technologies Pvt Ltd" --run --docx

Builds a realistic single-level MAKE scenario for manual MRP testing:
- 1 finished good VGS-MRP750-FG with a 5-component BOM (VGS-MRP750-RM-A..RM-E).
- A confirmed sales order for 750 finished units.
- Per-component BUY planning rules (safety stock, MOQ, order multiple, lead time).
- Opening stock incl. reserved and quarantine quantities, plus open POs.

Idempotent: re-running reuses records and never double-posts stock. Uses only
existing models and the inventory ledger service - no MRP logic is changed.

--run executes the MRP run and prints expected vs actual engine output.
--docx writes a Word manual-demo guide to media/demo_guides/ and prints its path.
"""
import datetime
import os
from decimal import Decimal, ROUND_CEILING

from django.conf import settings
from django.core.management.base import BaseCommand
from django.db import transaction
from django.utils import timezone

from core.models import (
    Tenant, Site, Location, Product, Supplier, Customer,
    BillOfMaterials, BillOfMaterialsLine, ItemSitePlanning,
    CustomerOrder, CustomerOrderLine, MRPRun, MRPPlannedOrder,
    PurchaseOrder, PurchaseOrderLine, InventoryBalance,
    WorkCentre, RoutingHeader, RoutingOperation,
)

SEED_REF = "MRP_750_SEED"
ZERO = Decimal("0.00")

PREFIX = "VGS-MRP750"            # SKU / document naming convention
FG_CODE = "FG"
FG_SKU = f"{PREFIX}-{FG_CODE}"   # VGS-MRP750-FG
SO_NUMBER = f"{PREFIX}-SO"
ROUTING_CODE = f"{PREFIX}-RT"
WORK_CENTRE_CODE = PREFIX
WORK_CENTRE_NAME = f"{PREFIX} Assembly"
DOCX_FILENAME = "VGS_MRP_750_Demo_Manual_Guide.docx"

ORDER_QTY = Decimal("750")
FG_MAKE_LEAD = 3

# code, bom_qty_per_fg, safety, moq, multiple, lead_days,
# on_hand (nettable, main store), reserved, quarantine, open_po
COMPONENTS = [
    ("RM-A", "2",   "500",  "500",  "100", 10, "900",  "100", "0",   "300"),
    ("RM-B", "1",   "300",  "100",  "50",   7, "1000", "200", "0",   "0"),
    ("RM-C", "4",   "1000", "1000", "500", 14, "2000", "0",   "200", "500"),
    ("RM-D", "0.5", "200",  "500",  "50",   7, "100",  "0",   "0",   "0"),
    ("RM-E", "1",   "250",  "250",  "50",   5, "900",  "0",   "0",   "0"),
]


def _sku(code):
    return f"{PREFIX}-{code}"


def _round_up_to_multiple(qty, multiple):
    if multiple <= ZERO:
        return qty
    lots = (qty / multiple).to_integral_value(rounding=ROUND_CEILING)
    return multiple * lots


def _classic_planned(net, moq, multiple):
    """Classic single-level lot sizing: max(net,0) -> MOQ -> order multiple."""
    qty = net if net > ZERO else ZERO
    if qty <= ZERO:
        return ZERO
    if moq > ZERO and qty < moq:
        qty = moq
    return _round_up_to_multiple(qty, multiple)


class Command(BaseCommand):
    help = "Seed (and optionally run + document) the VGS-MRP750 5-component MRP test scenario."

    def add_arguments(self, parser):
        parser.add_argument("--tenant", required=True, help="Tenant name (created if missing).")
        parser.add_argument("--run", action="store_true", help="Execute the MRP run and print actual vs expected.")
        parser.add_argument("--docx", action="store_true",
                            help="Write the Word manual-demo guide and print its path.")

    @transaction.atomic
    def handle(self, *args, **opts):
        tenant = self._tenant(opts["tenant"])
        self.today = timezone.localdate()
        # Sales order required well in the future so every release date is in the
        # future (no past-due noise): SO date - FG lead - longest component lead.
        self.so_required = self.today + datetime.timedelta(days=30)

        site = self._site(tenant)
        locs = self._locations(tenant, site)
        suppliers = self._suppliers(tenant)
        products = self._products(tenant)
        self._bom(tenant, products)
        self._routing(tenant, site, products)
        self._planning(tenant, site, products, suppliers)
        self._stock(tenant, products, locs)
        self._open_pos(tenant, site, products, suppliers, locs)
        self._sales(tenant, site, products)
        run = self._run(tenant, site)

        self._print_expected()
        self.stdout.write(self.style.SUCCESS(
            f"\nSeeded into '{tenant.name}'. MRP run: {run.run_number} (status {run.status}).\n"
            f"Open Planning -> MRP Runs -> {run.run_number} and click Run MRP, or re-run this "
            f"command with --run."))

        if opts["run"]:
            self._run_and_compare(run)

        if opts["docx"]:
            # The guide's Section 4 needs actual planned quantities, so ensure the
            # run has been executed before writing the document.
            if not MRPPlannedOrder.objects.filter(mrp_run=run).exists():
                from core.services.mrp.engine import run_mrp
                run_mrp(run)
            run.refresh_from_db()
            path = self._write_guide(tenant, run)
            self.stdout.write(self.style.SUCCESS(f"\nManual demo guide written to: {path}"))

    # ------------------------------------------------------------------ #
    def _tenant(self, ref):
        t = Tenant.objects.filter(name=ref).first()
        if t is None and str(ref).isdigit():
            t = Tenant.objects.filter(id=int(ref)).first()
        if t is None:
            t = Tenant.objects.create(name=ref)  # signal seeds accounts + default site/location
        return t

    def _site(self, tenant):
        site = Site.objects.filter(tenant=tenant).order_by("id").first()
        if site is None:
            site = Site.objects.create(tenant=tenant, name="Main Site",
                                       site_type=Site.Type.OPERATING_SITE, is_default=True)
        return site

    def _locations(self, tenant, site):
        def loc(name, ltype):
            return Location.objects.get_or_create(
                tenant=tenant, site=site, name=name,
                defaults={"type": ltype, "holds_stock": True, "is_active": True})[0]
        store = (Location.objects.filter(tenant=tenant, site=site, holds_stock=True, type="WAREHOUSE")
                 .order_by("id").first() or loc("Main Location", "WAREHOUSE"))
        return {"store": store, "quarantine": loc("Quarantine", "QUARANTINE")}

    def _suppliers(self, tenant):
        return {code: Supplier.objects.get_or_create(tenant=tenant, name=f"{PREFIX} Supplier {code}")[0]
                for code, *_rest in COMPONENTS}

    def _products(self, tenant):
        out = {}
        out[FG_CODE] = Product.objects.get_or_create(
            tenant=tenant, sku=FG_SKU,
            defaults={"name": f"{PREFIX} finished good", "product_type": "FINISHED_GOOD",
                      "standard_cost": Decimal("50.00"), "cost_method": "AVERAGE", "is_active": True})[0]
        for code, *_rest in COMPONENTS:
            out[code] = Product.objects.get_or_create(
                tenant=tenant, sku=_sku(code),
                defaults={"name": f"{PREFIX} raw material {code}", "product_type": "RAW_MATERIAL",
                          "standard_cost": Decimal("2.00"), "cost_method": "AVERAGE", "is_active": True})[0]
        return out

    def _bom(self, tenant, p):
        bom = BillOfMaterials.objects.get_or_create(
            tenant=tenant, product=p[FG_CODE], name="Default BOM",
            defaults={"output_qty": Decimal("1"), "is_active": True})[0]
        for i, (code, qty, *_rest) in enumerate(COMPONENTS, start=1):
            BillOfMaterialsLine.objects.get_or_create(
                bom=bom, component=p[code],
                defaults={"line_no": i * 10, "qty": Decimal(qty),
                          "scrap_percent": ZERO, "fixed_qty": ZERO})

    def _routing(self, tenant, site, p):
        # A simple single-operation routing so the MAKE item has an active routing
        # (no MISSING_ROUTING). No shop calendar / finite scheduling - keeps the
        # focus on material planning.
        wc = WorkCentre.objects.get_or_create(
            tenant=tenant, code=WORK_CENTRE_CODE,
            defaults={"site": site, "name": WORK_CENTRE_NAME, "capacity_hours_per_day": Decimal("8"),
                      "efficiency_percent": Decimal("100"), "finite_capacity_enabled": False,
                      "is_active": True})[0]
        routing, created = RoutingHeader.objects.get_or_create(
            tenant=tenant, routing_code=ROUTING_CODE,
            defaults={"product": p[FG_CODE], "site": None, "status": "ACTIVE", "is_default": True})
        if created:
            # Small per-unit time so 750 units fit one shift (no capacity overload).
            RoutingOperation.objects.create(
                routing=routing, operation_sequence=10, operation_name="Assemble",
                work_centre=wc, setup_minutes=Decimal("10"), run_minutes_per_unit=Decimal("0.1"))

    def _planning(self, tenant, site, p, suppliers):
        # FG: MAKE, driven by the sales order.
        ItemSitePlanning.objects.get_or_create(
            tenant=tenant, product=p[FG_CODE], site=site,
            defaults=dict(source_type="MAKE", mrp_enabled=True, is_active=True,
                          include_sales_orders=True, include_forecast=False, include_safety_stock=True,
                          safety_stock_qty=ZERO, lead_time_days=FG_MAKE_LEAD,
                          lot_sizing_method="LOT_FOR_LOT"))
        for code, _qty, safety, moq, mult, lead, *_rest in COMPONENTS:
            ItemSitePlanning.objects.get_or_create(
                tenant=tenant, product=p[code], site=site,
                defaults=dict(source_type="BUY", mrp_enabled=True, is_active=True,
                              include_sales_orders=True, include_forecast=False, include_safety_stock=True,
                              default_supplier=suppliers[code],
                              safety_stock_qty=Decimal(safety), min_order_qty=Decimal(moq),
                              order_multiple=Decimal(mult), lead_time_days=lead,
                              lot_sizing_method="LOT_FOR_LOT"))

    def _receive(self, tenant, product, location, qty, ref):
        if Decimal(qty) <= ZERO:
            return
        if InventoryMovement_exists(tenant, product, location, ref):
            return
        from core.services.inventory import apply_movement
        apply_movement(tenant=tenant, product=product, location=location, movement_type="RECEIVE",
                       qty_delta=Decimal(qty), ref_type=SEED_REF, ref_id=ref,
                       unit_cost=Decimal("2.00"))

    def _stock(self, tenant, p, locs):
        for code, _qty, _safety, _moq, _mult, _lead, on_hand, reserved, quar, _po in COMPONENTS:
            self._receive(tenant, p[code], locs["store"], on_hand, f"{code}-onhand")
            self._receive(tenant, p[code], locs["quarantine"], quar, f"{code}-quar")
            # Reserved is set on the balance directly (idempotent: assignment, not increment).
            if Decimal(reserved) > ZERO:
                bal = InventoryBalance.objects.filter(
                    tenant=tenant, product=p[code], location=locs["store"]).first()
                if bal and bal.reserved != Decimal(reserved):
                    bal.reserved = Decimal(reserved)
                    bal.save(update_fields=["reserved"])

    def _open_pos(self, tenant, site, p, suppliers, locs):
        expected = self.today + datetime.timedelta(days=5)  # due before production
        for code, _qty, _safety, _moq, _mult, _lead, _oh, _res, _quar, po_qty in COMPONENTS:
            if Decimal(po_qty) <= ZERO:
                continue
            number = f"{PREFIX}-PO-{code}"
            if PurchaseOrder.objects.filter(tenant=tenant, po_number=number).exists():
                continue
            po = PurchaseOrder.objects.create(
                tenant=tenant, po_number=number, supplier=suppliers[code], site=site,
                receiving_location=locs["store"], status="APPROVED", is_current=True,
                expected_date=expected)
            PurchaseOrderLine.objects.create(
                po=po, product=p[code], ordered_qty=Decimal(po_qty), received_qty=ZERO,
                unit_cost=Decimal("2.00"))

    def _sales(self, tenant, site, p):
        cust = Customer.objects.get_or_create(tenant=tenant, name=f"{PREFIX} Demo Customer")[0]
        if not CustomerOrder.objects.filter(tenant=tenant, order_number=SO_NUMBER).exists():
            o = CustomerOrder.objects.create(
                tenant=tenant, customer=cust, site=site, order_number=SO_NUMBER,
                status="CONFIRMED", order_date=self.so_required)
            CustomerOrderLine.objects.create(order=o, product=p[FG_CODE], qty=ORDER_QTY)

    def _run(self, tenant, site):
        from core.services.mrp import next_run_number
        run = MRPRun.objects.filter(tenant=tenant, run_number__startswith=f"{PREFIX}-").first()
        if run is None:
            run = MRPRun.objects.create(
                tenant=tenant, run_number=f"{PREFIX}-" + next_run_number(tenant)[-6:],
                site_scope=site, planning_start_date=self.today,
                planning_end_date=self.today + datetime.timedelta(days=180),
                include_sales_orders=True, include_forecast=False, include_safety_stock=True,
                include_transfers=True, status="DRAFT")
        return run

    # ------------------------------------------------------------------ #
    # Expected / actual calculation (shared by the console table and the guide)
    # ------------------------------------------------------------------ #
    def _expected_rows(self):
        """List of dicts with the full expected calculation for each component."""
        rows = []
        for code, qty, safety, moq, mult, lead, oh, res, quar, po in COMPONENTS:
            gross = ORDER_QTY * Decimal(qty)
            nettable = Decimal(oh) - Decimal(res)              # quarantine fully excluded
            usable = nettable + Decimal(po)
            net = gross + Decimal(safety) - usable
            planned = _classic_planned(net, Decimal(moq), Decimal(mult))
            rows.append(dict(
                sku=_sku(code), qty_per_fg=qty, gross=gross, safety=Decimal(safety),
                on_hand=Decimal(oh), reserved=Decimal(res), quarantine=Decimal(quar),
                open_po=Decimal(po), usable=usable, net=net, moq=Decimal(moq),
                multiple=Decimal(mult), planned=planned, lead=lead))
        return rows

    def _actual_planned(self, run):
        """(sku, source_type) -> total planned quantity from the run."""
        out = {}
        for po in MRPPlannedOrder.objects.filter(mrp_run=run).select_related("product"):
            key = (po.product.sku, po.source_type)
            out[key] = out.get(key, ZERO) + (po.quantity or ZERO)
        return out

    def _print_expected(self):
        fg_release = self.so_required - datetime.timedelta(days=FG_MAKE_LEAD)
        self.stdout.write("\nExpected (classic single-level netting incl. safety stock):")
        self.stdout.write(
            "  Component         Q/FG  Gross  Safety  OnHand  Resv  Quar  OpenPO  Usable  Net   MOQ   Mult  Planned")
        for r in self._expected_rows():
            self.stdout.write(
                f"  {r['sku']:16} {r['qty_per_fg']:>4}  {r['gross']:>5}  {r['safety']:>5}  "
                f"{r['on_hand']:>5}  {r['reserved']:>4}  {r['quarantine']:>4}  {r['open_po']:>5}  "
                f"{r['usable']:>6}  {r['net']:>4}  {r['moq']:>4}  {r['multiple']:>4}  {r['planned']:>7}")
        self.stdout.write(f"\n  FG {FG_SKU}: MAKE {ORDER_QTY}, required {self.so_required}, release {fg_release}.")

    def _run_and_compare(self, run):
        from core.services.mrp.engine import run_mrp
        from core.models import MRPDemand, MRPException
        run = MRPRun.objects.get(pk=run.pk)
        run_mrp(run)
        run.refresh_from_db()
        self.stdout.write(self.style.WARNING(f"\nActual engine output (run {run.run_number}, status {run.status}):"))

        actual = self._actual_planned(run)
        self.stdout.write("  Planned orders by product:")
        self.stdout.write(f"    {FG_SKU:16} MAKE qty {actual.get((FG_SKU, 'MAKE'), ZERO)}")
        for code, *_rest in COMPONENTS:
            self.stdout.write(f"    {_sku(code):16} BUY  qty {actual.get((_sku(code), 'BUY'), ZERO)}")

        dem = (MRPDemand.objects.filter(mrp_run=run, demand_type="WORK_ORDER_COMPONENT")
               .select_related("product").order_by("product__sku"))
        self.stdout.write(f"\n  WORK_ORDER_COMPONENT demand rows: {dem.count()}")
        for d in dem:
            self.stdout.write(f"    {d.product.sku:16} qty {d.quantity:>7} required {d.required_date}")

        excs = MRPException.objects.filter(mrp_run=run)
        seen = {}
        for e in excs:
            seen[e.exception_code] = seen.get(e.exception_code, 0) + 1
        self.stdout.write(f"\n  Exceptions: {excs.count()}")
        for code, n in sorted(seen.items()):
            self.stdout.write(f"    {code} x{n}")

    # ------------------------------------------------------------------ #
    # Word (.docx) manual demo guide
    # ------------------------------------------------------------------ #
    def _guide_path(self):
        out_dir = os.path.join(settings.MEDIA_ROOT, "demo_guides")
        os.makedirs(out_dir, exist_ok=True)
        return os.path.join(out_dir, DOCX_FILENAME)

    def _write_guide(self, tenant, run):
        try:
            from docx import Document  # noqa: F401
        except Exception:
            return self._write_markdown_fallback(tenant, run)
        return self._write_docx(tenant, run)

    def _write_docx(self, tenant, run):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        actual = self._actual_planned(run)
        fg_actual = actual.get((FG_SKU, "MAKE"), ZERO)
        rows = self._expected_rows()
        path = self._guide_path()

        doc = Document()
        title = doc.add_heading("VGS MRP 750 Demo Manual Guide", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(f"Generated for {tenant.name} - MRP run {run.run_number}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in sub.runs:
            r.italic = True

        # Section 1: Demo Purpose
        doc.add_heading("1. Demo Purpose", level=1)
        doc.add_paragraph(
            "This demo walks through the full material-planning process end to end:")
        for step in [
            "Sales Order for 750 finished goods",
            "BOM explosion into component demand",
            "Material availability check",
            "Reserved and quarantine stock excluded from usable supply",
            "Open purchase-order supply counted",
            "Safety stock included in the requirement",
            "MOQ and order multiple applied",
            "Planned BUY and MAKE orders created",
            "Pegging that links demand to supply",
            "Conversion to Purchase Requisitions / Purchase Orders and a Work Order",
            "Validation in the MRP reports and the planner dashboard",
        ]:
            doc.add_paragraph(step, style="List Bullet")

        # Section 2: Demo Company
        doc.add_heading("2. Demo Company", level=1)
        doc.add_paragraph(f"Company: {tenant.name}")

        # Section 3: Demo Records Created
        doc.add_heading("3. Demo Records Created", level=1)
        self._docx_kv(doc, "Products", [
            f"{FG_SKU} (finished good)",
            *[_sku(c[0]) for c in COMPONENTS],
        ])
        self._docx_kv(doc, "BOM", [f"{FG_SKU} uses 5 raw materials (RM-A..RM-E)"])
        self._docx_kv(doc, "Sales Order", [f"{SO_NUMBER} - quantity {ORDER_QTY}"])
        self._docx_kv(doc, "Open POs", [
            f"{PREFIX}-PO-RM-A quantity 300",
            f"{PREFIX}-PO-RM-C quantity 500",
        ])
        self._docx_kv(doc, "Routing", [f"{ROUTING_CODE}", f"Work Centre: {WORK_CENTRE_NAME}"])
        self._docx_kv(doc, "MRP Run", [f"{run.run_number} (status {run.get_status_display()})"])

        # Section 4: Expected MRP Calculation
        doc.add_heading("4. Expected MRP Calculation", level=1)
        headers = ["Component", "Qty/FG", "Gross", "Safety", "On Hand", "Reserved", "Quarantine",
                   "Open PO", "Usable", "Net", "MOQ", "Order Mult", "Expected Qty", "Actual Qty", "Status"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(8)
        # FG row
        self._docx_row(table, [FG_SKU, "-", str(ORDER_QTY), "0", "0", "0", "0", "0", "0",
                               str(ORDER_QTY), "-", "-", str(ORDER_QTY), str(fg_actual),
                               "OK" if fg_actual == ORDER_QTY else "CHECK"])
        # Component rows
        for r in rows:
            act = actual.get((r["sku"], "BUY"), ZERO)
            self._docx_row(table, [
                r["sku"], r["qty_per_fg"], str(r["gross"]), str(r["safety"]), str(r["on_hand"]),
                str(r["reserved"]), str(r["quarantine"]), str(r["open_po"]), str(r["usable"]),
                str(r["net"]), str(r["moq"]), str(r["multiple"]), str(r["planned"]), str(act),
                "OK" if act == r["planned"] else "CHECK"])
        doc.add_paragraph(
            "Note: RM-D is planned with exactly one BUY order of 500 (no duplicate). "
            "The 200 quarantine units of RM-C are excluded from usable supply.").italic = True

        # Section 5: Manual UI Steps
        doc.add_heading("5. Manual UI Steps", level=1)
        for line in self._ui_steps(run):
            if line.startswith("  - "):
                doc.add_paragraph(line[4:], style="List Bullet 2")
            else:
                doc.add_paragraph(line, style="List Number")

        # Section 6: One-Click Conversion Command
        doc.add_heading("6. One-Click Conversion Command", level=1)
        self._docx_code(doc, f'python manage.py demo_mrp_750_convert --tenant "{tenant.name}"')
        for line in [
            "Converts BUY planned orders to Purchase Requisitions.",
            "Converts the MAKE planned order to a Work Order.",
            "Safe and idempotent - re-running does not duplicate PRs or Work Orders.",
        ]:
            doc.add_paragraph(line, style="List Bullet")
        doc.add_paragraph("Optional - convert BUY orders to draft Purchase Orders instead:")
        self._docx_code(doc, f'python manage.py demo_mrp_750_convert --tenant "{tenant.name}" --buy-as po')
        doc.add_paragraph("Optional - execute the work order (issues materials, posts inventory):")
        self._docx_code(doc, f'python manage.py demo_mrp_750_convert --tenant "{tenant.name}" --execute-work-order')
        warn = doc.add_paragraph(
            "Warning: do not use --execute-work-order during a clean demo unless you "
            "intentionally want inventory movements posted.")
        for r in warn.runs:
            r.bold = True

        # Section 7: Demo Talking Points
        doc.add_heading("7. Demo Talking Points", level=1)
        for line in [
            "The system does not just check stock on hand - it checks usable stock.",
            "Reserved stock is excluded.",
            "Quarantine stock is excluded.",
            "Open POs are counted if due on time.",
            "Safety stock is maintained.",
            "Supplier MOQ and order multiples are respected.",
            "Lead times decide release dates.",
            "Pegging explains why each planned order exists.",
            "The planner can convert suggestions into real PRs, POs, and Work Orders.",
        ]:
            doc.add_paragraph(line, style="List Bullet")

        # Section 8: Reset / Re-run Notes
        doc.add_heading("8. Reset / Re-run Notes", level=1)
        for line in [
            "The seed command is idempotent - re-running reuses and updates demo records.",
            "Re-running the seed never double-posts opening stock.",
            "If the Work Order execution command was run, inventory movements may be posted.",
            "For a clean demo, reseed into a fresh tenant or reset the demo records.",
        ]:
            doc.add_paragraph(line, style="List Bullet")

        doc.save(path)
        return path

    def _docx_kv(self, doc, label, items):
        p = doc.add_paragraph()
        p.add_run(f"{label}:").bold = True
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    def _docx_row(self, table, values):
        from docx.shared import Pt
        cells = table.add_row().cells
        for i, v in enumerate(values):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    def _docx_code(self, doc, text):
        from docx.shared import Pt
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(9)

    def _ui_steps(self, run):
        return [
            "Login to the ERP.",
            "Switch company to the demo tenant.",
            "Go to Products and search VGS-MRP750.",
            f"Open {FG_SKU}.",
            f"Go to BOMs and open the BOM for {FG_SKU}.",
            "Confirm the 5 component lines.",
            "Go to Item Planning.",
            "Confirm planning profiles for FG and RM-A to RM-E.",
            "Go to Inventory.",
            "Confirm nettable stock and quarantine stock.",
            "Go to Purchase Orders.",
            "Confirm open POs for RM-A and RM-C.",
            "Go to Sales Orders.",
            "Confirm the 750-unit sales order.",
            "Go to Planning -> MRP Runs.",
            f"Open the generated MRP run {run.run_number}.",
            "Click Run MRP if not already run.",
            "Open the Planner Workbench.",
            "Check the Demand section.",
            "Confirm sales demand for FG quantity 750.",
            "Confirm component demands:",
            "  - RM-A 1500",
            "  - RM-B 750",
            "  - RM-C 3000",
            "  - RM-D 375",
            "  - RM-E 750",
            "Check Planned Orders. Confirm:",
            "  - FG MAKE 750",
            "  - RM-A BUY 900",
            "  - RM-B BUY 250",
            "  - RM-C BUY 1500",
            "  - RM-D BUY 500",
            "  - RM-E BUY 250",
            "Open Pegging.",
            "Confirm sales order -> MAKE planned order -> component demand -> BUY planned orders.",
            "Open Exceptions and confirm only genuine exceptions appear.",
            "Convert BUY planned orders to Purchase Requisitions.",
            "Convert the MAKE planned order to a Work Order.",
            "Go to Requisitions and confirm PRs were created.",
            "Go to Work Orders and confirm the Work Order and material lines.",
            "Go to MRP Reports.",
            "Validate Planned Orders, Demand vs Supply, Pegging, and Exceptions.",
            "Go to the Planner Dashboard and validate the dashboard cards.",
        ]

    def _write_markdown_fallback(self, tenant, run):
        """Used only if python-docx is unavailable: a Markdown guide instead."""
        path = self._guide_path().replace(".docx", ".md")
        actual = self._actual_planned(run)
        lines = [f"# VGS MRP 750 Demo Manual Guide", "",
                 f"Company: {tenant.name}", f"MRP run: {run.run_number}", "",
                 "## Expected MRP Calculation", ""]
        for r in self._expected_rows():
            act = actual.get((r["sku"], "BUY"), ZERO)
            lines.append(f"- {r['sku']}: expected {r['planned']}, actual {act}")
        lines += ["", "## Manual UI Steps", ""]
        n = 1
        for s in self._ui_steps(run):
            if s.startswith("  - "):
                lines.append(f"    {s.strip()}")
            else:
                lines.append(f"{n}. {s}")
                n += 1
        with open(path, "w", encoding="utf-8") as fh:
            fh.write("\n".join(lines))
        return path


def InventoryMovement_exists(tenant, product, location, ref):
    from core.models import InventoryMovement
    return InventoryMovement.objects.filter(
        tenant=tenant, ref_type=SEED_REF, ref_id=ref, product=product, location=location).exists()
